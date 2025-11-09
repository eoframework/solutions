#!/usr/bin/env python3
"""
Analyze Word document to extract detailed formatting information.
Used for reverse-engineering templates.
"""

import argparse
from docx import Document
from docx.shared import Pt
from pathlib import Path


def rgb_to_hex(rgb_color):
    """Convert RGBColor to hex string."""
    if rgb_color is None:
        return None
    try:
        return f"#{rgb_color:06X}"
    except:
        return None


def analyze_paragraph(para, index):
    """Analyze a paragraph and return formatting details."""
    info = {
        'index': index,
        'text': para.text[:100] + ('...' if len(para.text) > 100 else ''),
        'style': para.style.name if para.style else 'None',
        'alignment': str(para.alignment) if para.alignment else 'None',
    }

    # Get first run's formatting if available
    if para.runs:
        first_run = para.runs[0]
        info['font_name'] = first_run.font.name
        info['font_size'] = str(first_run.font.size) if first_run.font.size else 'None'
        info['bold'] = first_run.font.bold
        info['italic'] = first_run.font.italic

        # Try to get color
        if first_run.font.color and first_run.font.color.rgb:
            info['color'] = rgb_to_hex(first_run.font.color.rgb)
        else:
            info['color'] = 'None'

    return info


def analyze_table(table, table_num):
    """Analyze a table and return structure details."""
    info = {
        'table_num': table_num,
        'rows': len(table.rows),
        'cols': len(table.columns),
        'style': table.style.name if table.style else 'None'
    }

    # Analyze first row (usually header)
    if table.rows:
        first_row = table.rows[0]
        header_cells = [cell.text[:30] for cell in first_row.cells[:5]]  # First 5 cells
        info['first_row'] = header_cells

    return info


def analyze_document(docx_path):
    """Analyze a Word document and print detailed formatting information."""
    print("=" * 80)
    print(f"ANALYZING WORD DOCUMENT: {docx_path.name}")
    print("=" * 80)

    doc = Document(docx_path)

    # Analyze styles
    print("\n📐 DOCUMENT STYLES:")
    print("-" * 80)

    key_styles = ['Normal', 'Title', 'Heading 1', 'Heading 2', 'Heading 3']
    for style_name in key_styles:
        try:
            style = doc.styles[style_name]
            font = style.font

            print(f"\n{style_name}:")
            print(f"  Font: {font.name if font.name else 'Not set'}")
            print(f"  Size: {font.size if font.size else 'Not set'}")
            print(f"  Bold: {font.bold if font.bold is not None else 'Not set'}")
            print(f"  Color: {rgb_to_hex(font.color.rgb) if font.color and font.color.rgb else 'Not set'}")
        except KeyError:
            print(f"\n{style_name}: NOT FOUND")

    # Analyze paragraphs
    print("\n\n📝 PARAGRAPHS:")
    print("-" * 80)

    for idx, para in enumerate(doc.paragraphs[:50]):  # First 50 paragraphs
        if para.text.strip():  # Only non-empty paragraphs
            info = analyze_paragraph(para, idx)
            print(f"\nPara {info['index']}: {info['style']}")
            print(f"  Text: {info['text']}")
            print(f"  Font: {info.get('font_name', 'N/A')} {info.get('font_size', 'N/A')}")
            print(f"  Bold: {info.get('bold', 'N/A')}, Italic: {info.get('italic', 'N/A')}")
            print(f"  Color: {info.get('color', 'N/A')}")
            print(f"  Alignment: {info['alignment']}")

    # Analyze tables
    print("\n\n📊 TABLES:")
    print("-" * 80)

    for idx, table in enumerate(doc.tables):
        info = analyze_table(table, idx + 1)
        print(f"\nTable {info['table_num']}:")
        print(f"  Dimensions: {info['rows']} rows × {info['cols']} cols")
        print(f"  Style: {info['style']}")
        print(f"  First row: {info.get('first_row', [])}")

        # Analyze cell formatting (first cell)
        if table.rows and table.rows[0].cells:
            first_cell = table.rows[0].cells[0]
            if first_cell.paragraphs:
                cell_para = first_cell.paragraphs[0]
                if cell_para.runs:
                    run = cell_para.runs[0]
                    print(f"  Header cell font: {run.font.name} {run.font.size}")
                    print(f"  Header cell bold: {run.font.bold}")

    # Analyze sections (for headers/footers)
    print("\n\n📄 SECTIONS:")
    print("-" * 80)

    for idx, section in enumerate(doc.sections):
        print(f"\nSection {idx + 1}:")
        print(f"  Page width: {section.page_width}")
        print(f"  Page height: {section.page_height}")
        print(f"  Left margin: {section.left_margin}")
        print(f"  Right margin: {section.right_margin}")
        print(f"  Top margin: {section.top_margin}")
        print(f"  Bottom margin: {section.bottom_margin}")

        # Analyze footer
        if section.footer:
            footer = section.footer
            print(f"  Footer paragraphs: {len(footer.paragraphs)}")
            if footer.paragraphs:
                footer_para = footer.paragraphs[0]
                print(f"  Footer text: {footer_para.text[:100]}")

    print("\n" + "=" * 80)
    print("ANALYSIS COMPLETE")
    print("=" * 80)


if __name__ == '__main__':
    parser = argparse.ArgumentParser(
        description='Analyze Word document formatting for reverse engineering.'
    )
    parser.add_argument(
        'docx_file',
        type=Path,
        help='Path to Word document (.docx) to analyze'
    )

    args = parser.parse_args()

    if not args.docx_file.exists():
        print(f"Error: File not found: {args.docx_file}")
        exit(1)

    analyze_document(args.docx_file)

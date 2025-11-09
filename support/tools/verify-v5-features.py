#!/usr/bin/env python3
"""
Verify v5 template features (padding and 12pt fonts).
"""

from docx import Document
from pathlib import Path
import sys

docx_path = Path(sys.argv[1])
doc = Document(docx_path)

print("=" * 80)
print("TEMPLATE V5 VERIFICATION")
print("=" * 80)

# 1. Check metadata table padding
print("\n1. METADATA TABLE PADDING:")
print("-" * 80)

if doc.tables:
    metadata_table = doc.tables[0]
    print(f"Table dimensions: {len(metadata_table.rows)} rows × {len(metadata_table.columns)} cols")

    # Check first row for padding
    first_row = metadata_table.rows[0]
    for cell_idx, cell in enumerate(first_row.cells):
        if cell.paragraphs:
            para = cell.paragraphs[0]
            space_before = para.paragraph_format.space_before
            space_after = para.paragraph_format.space_after
            print(f"\n  Cell {cell_idx + 1} ('{cell.text[:30]}')")
            print(f"    Space before: {space_before}")
            print(f"    Space after: {space_after}")
            if space_before and space_after:
                print(f"    ✅ Has padding")
            else:
                print(f"    ❌ No padding")

# 2. Check all font sizes
print("\n\n2. FONT SIZES VERIFICATION:")
print("-" * 80)

font_sizes = {}

# Check paragraphs
for idx, para in enumerate(doc.paragraphs[:100]):
    for run in para.runs:
        if run.font.size:
            size_pt = run.font.size.pt
            if size_pt not in font_sizes:
                font_sizes[size_pt] = []
            font_sizes[size_pt].append(f"Para {idx}: '{para.text[:40]}'")

# Check tables
for table_idx, table in enumerate(doc.tables):
    for row_idx, row in enumerate(table.rows):
        for cell_idx, cell in enumerate(row.cells):
            for para in cell.paragraphs:
                for run in para.runs:
                    if run.font.size:
                        size_pt = run.font.size.pt
                        if size_pt not in font_sizes:
                            font_sizes[size_pt] = []
                        font_sizes[size_pt].append(f"Table {table_idx} Row {row_idx} Cell {cell_idx}")

# Display summary
print("\nFont sizes found:")
for size in sorted(font_sizes.keys()):
    count = len(font_sizes[size])
    print(f"\n  {size}pt: {count} occurrences")
    if size != 12 and size != 16 and size != 28:  # Allow 16pt for headings, 28pt for title
        print(f"    ⚠️  Non-standard size! Examples:")
        for example in font_sizes[size][:3]:
            print(f"      - {example}")

# 3. Check footer font sizes
print("\n\n3. FOOTER FONT SIZES:")
print("-" * 80)

for section_idx, section in enumerate(doc.sections):
    if section.footer.paragraphs:
        print(f"\nSection {section_idx + 1} footer:")
        for para in section.footer.paragraphs:
            for run_idx, run in enumerate(para.runs):
                if run.text.strip():
                    size = run.font.size.pt if run.font.size else 'None'
                    print(f"  Run {run_idx}: '{run.text[:30]}' - {size}pt")

print("\n" + "=" * 80)
print("VERIFICATION COMPLETE")
print("=" * 80)

#!/usr/bin/env python3
"""
Detailed analyzer for specific Word document features.
Focus on: footer, metadata table, cover page, TOC line spacing.
"""

import argparse
from docx import Document
from docx.shared import Pt, Inches
from docx.oxml.ns import qn
from pathlib import Path


def pt_to_inches(pt_value):
    """Convert EMUs to inches."""
    if pt_value:
        return pt_value / 914400
    return None


def analyze_footer_details(doc):
    """Analyze footer structure and content."""
    print("\n" + "=" * 80)
    print("FOOTER ANALYSIS")
    print("=" * 80)

    for idx, section in enumerate(doc.sections):
        print(f"\nSection {idx + 1} Footer:")
        footer = section.footer

        if not footer.paragraphs:
            print("  No footer paragraphs found")
            continue

        for para_idx, para in enumerate(footer.paragraphs):
            print(f"\n  Paragraph {para_idx + 1}:")
            print(f"    Text: '{para.text}'")
            print(f"    Alignment: {para.alignment}")

            # Check tab stops
            tab_stops = para.paragraph_format.tab_stops
            if tab_stops:
                print(f"    Tab stops: {len(list(tab_stops))} found")
                for tab in tab_stops:
                    print(f"      - Position: {pt_to_inches(tab.position):.2f}\" Alignment: {tab.alignment}")

            # Check runs
            print(f"    Runs: {len(para.runs)}")
            for run_idx, run in enumerate(para.runs):
                if run.text.strip():
                    print(f"      Run {run_idx + 1}: '{run.text}' Font: {run.font.name} Size: {run.font.size}")

            # Check for images
            for run in para.runs:
                if hasattr(run, '_element'):
                    drawings = run._element.findall('.//{http://schemas.openxmlformats.org/wordprocessingml/2006/main}drawing')
                    if drawings:
                        print(f"    Contains {len(drawings)} image(s)")


def analyze_metadata_table_details(doc):
    """Analyze metadata table (first table) in detail."""
    print("\n" + "=" * 80)
    print("METADATA TABLE ANALYSIS")
    print("=" * 80)

    if not doc.tables:
        print("No tables found")
        return

    table = doc.tables[0]
    print(f"\nTable dimensions: {len(table.rows)} rows × {len(table.columns)} cols")

    # Column widths
    print("\nColumn widths:")
    for col_idx, col in enumerate(table.columns):
        width_inches = pt_to_inches(col.width)
        print(f"  Column {col_idx + 1}: {width_inches:.3f}\" ({col.width} EMUs)")

    # Analyze first row in detail
    print("\nFirst row cells:")
    if table.rows:
        row = table.rows[0]
        for cell_idx, cell in enumerate(row.cells):
            print(f"\n  Cell {cell_idx + 1}:")
            print(f"    Text: '{cell.text}'")
            print(f"    Width: {pt_to_inches(cell.width):.3f}\"")
            print(f"    Vertical alignment: {cell.vertical_alignment}")

            # Check cell background (shading)
            tc = cell._element
            tcPr = tc.find(qn('w:tcPr'))
            if tcPr is not None:
                shd = tcPr.find(qn('w:shd'))
                if shd is not None:
                    fill_color = shd.get(qn('w:fill'))
                    print(f"    Background color: {fill_color}")

                # Check borders
                tcBorders = tcPr.find(qn('w:tcBorders'))
                if tcBorders is not None:
                    print("    Borders:")
                    for edge in ['top', 'bottom', 'left', 'right']:
                        border = tcBorders.find(qn(f'w:{edge}'))
                        if border is not None:
                            color = border.get(qn('w:color'))
                            sz = border.get(qn('w:sz'))
                            val = border.get(qn('w:val'))
                            print(f"      {edge}: color={color}, size={sz}, type={val}")

            # Check font in cell
            if cell.paragraphs and cell.paragraphs[0].runs:
                run = cell.paragraphs[0].runs[0]
                print(f"    Font: {run.font.name} {run.font.size}")
                print(f"    Bold: {run.font.bold}")


def analyze_cover_page_details(doc):
    """Analyze cover page layout and formatting."""
    print("\n" + "=" * 80)
    print("COVER PAGE ANALYSIS (First 10 paragraphs)")
    print("=" * 80)

    for idx, para in enumerate(doc.paragraphs[:10]):
        print(f"\nParagraph {idx}:")
        print(f"  Text: '{para.text[:80]}'")
        print(f"  Style: {para.style.name if para.style else 'None'}")
        print(f"  Alignment: {para.alignment}")

        if para.runs:
            first_run = para.runs[0]
            print(f"  Font: {first_run.font.name} {first_run.font.size}")
            print(f"  Bold: {first_run.font.bold}, Italic: {first_run.font.italic}")

        # Check for images
        for run in para.runs:
            if hasattr(run, '_element'):
                drawings = run._element.findall('.//{http://schemas.openxmlformats.org/wordprocessingml/2006/main}drawing')
                blips = run._element.findall('.//{http://schemas.openxmlformats.org/drawingml/2006/main}blip')
                if drawings or blips:
                    # Try to get image dimensions
                    extents = run._element.findall('.//{http://schemas.openxmlformats.org/drawingml/2006/wordprocessingDrawing}extent')
                    if extents:
                        for extent in extents:
                            cx = extent.get('cx')
                            cy = extent.get('cy')
                            if cx and cy:
                                width_inches = int(cx) / 914400
                                height_inches = int(cy) / 914400
                                print(f"  Image: {width_inches:.2f}\" × {height_inches:.2f}\"")


def analyze_toc_line_spacing(doc):
    """Analyze Table of Contents line spacing."""
    print("\n" + "=" * 80)
    print("TABLE OF CONTENTS LINE SPACING ANALYSIS")
    print("=" * 80)

    # Find TOC heading
    toc_found = False
    for idx, para in enumerate(doc.paragraphs):
        if 'Table of Contents' in para.text and para.style.name == 'Heading 1':
            toc_found = True
            print(f"\nTOC found at paragraph {idx}")

            # Analyze next 20 paragraphs (TOC entries)
            print("\nTOC entries line spacing:")
            for i in range(idx + 1, min(idx + 21, len(doc.paragraphs))):
                toc_para = doc.paragraphs[i]
                if toc_para.text.strip() and not toc_para.text.startswith('List of Figures'):
                    line_spacing = toc_para.paragraph_format.line_spacing
                    line_spacing_rule = toc_para.paragraph_format.line_spacing_rule
                    print(f"  Para {i}: '{toc_para.text[:60]}'")
                    print(f"    Line spacing: {line_spacing} Rule: {line_spacing_rule}")
                    print(f"    Font size: {toc_para.runs[0].font.size if toc_para.runs else 'N/A'}")
                elif 'List of Figures' in toc_para.text:
                    break
            break

    if not toc_found:
        print("Table of Contents not found")


def analyze_page_structure(doc):
    """Analyze page structure and breaks."""
    print("\n" + "=" * 80)
    print("PAGE STRUCTURE ANALYSIS")
    print("=" * 80)

    page_breaks = []
    for idx, para in enumerate(doc.paragraphs[:30]):
        # Check for page breaks
        for run in para.runs:
            if hasattr(run, '_element'):
                breaks = run._element.findall('.//{http://schemas.openxmlformats.org/wordprocessingml/2006/main}br')
                for br in breaks:
                    br_type = br.get(qn('w:type'))
                    if br_type == 'page':
                        page_breaks.append(idx)
                        print(f"\nPage break found after paragraph {idx}")
                        if idx > 0:
                            print(f"  Previous para text: '{doc.paragraphs[idx - 1].text[:60]}'")
                        if idx < len(doc.paragraphs) - 1:
                            print(f"  Next para text: '{doc.paragraphs[idx + 1].text[:60]}'")

    print(f"\nTotal page breaks found in first 30 paragraphs: {len(page_breaks)}")


if __name__ == '__main__':
    parser = argparse.ArgumentParser(
        description='Detailed analysis of Word document specific features.'
    )
    parser.add_argument(
        'docx_file',
        type=Path,
        help='Path to Word document (.docx) to analyze'
    )

    args = parser.parse_args()

    if not args.docx_file.exists():
        print(f"Error: File not found: {args.docx_file}")
        exit(1)

    doc = Document(args.docx_file)

    analyze_cover_page_details(doc)
    analyze_metadata_table_details(doc)
    analyze_toc_line_spacing(doc)
    analyze_page_structure(doc)
    analyze_footer_details(doc)

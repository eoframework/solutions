#!/usr/bin/env python3
"""
Check metadata table line spacing in detail.
"""

from docx import Document
from pathlib import Path
import sys

docx_path = Path(sys.argv[1])
doc = Document(docx_path)

print("=" * 80)
print("METADATA TABLE LINE SPACING CHECK")
print("=" * 80)

if not doc.tables:
    print("No tables found")
    exit(1)

table = doc.tables[0]
print(f"\nTable: {len(table.rows)} rows × {len(table.columns)} cols\n")

for row_idx, row in enumerate(table.rows):
    print(f"Row {row_idx + 1}:")
    for col_idx, cell in enumerate(row.cells):
        if cell.paragraphs:
            para = cell.paragraphs[0]
            line_spacing = para.paragraph_format.line_spacing
            line_spacing_rule = para.paragraph_format.line_spacing_rule
            print(f"  Cell {col_idx + 1}: '{cell.text[:30]}'")
            print(f"    Line spacing: {line_spacing} Rule: {line_spacing_rule}")

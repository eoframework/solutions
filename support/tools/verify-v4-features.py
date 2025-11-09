#!/usr/bin/env python3
"""
Verify v4 template features.
"""

from docx import Document
from pathlib import Path
import sys

docx_path = Path(sys.argv[1])
doc = Document(docx_path)

print("=" * 80)
print("TEMPLATE V4 VERIFICATION")
print("=" * 80)

# 1. Check TOC spacing
print("\n1. TABLE OF CONTENTS SPACING:")
print("-" * 80)

toc_found = False
for idx, para in enumerate(doc.paragraphs):
    if 'Table of Contents' in para.text and para.style.name == 'Heading 1':
        toc_found = True
        # Check next few TOC entries
        for i in range(idx + 2, min(idx + 6, len(doc.paragraphs))):
            toc_para = doc.paragraphs[i]
            if toc_para.text.strip() and '.' in toc_para.text:
                space_before = toc_para.paragraph_format.space_before
                space_after = toc_para.paragraph_format.space_after
                print(f"TOC entry: '{toc_para.text[:50]}'")
                print(f"  Space before: {space_before}")
                print(f"  Space after: {space_after}")
                break
        break

# 2. Check bullet point indents
print("\n2. BULLET POINT INDENTS:")
print("-" * 80)

for idx, para in enumerate(doc.paragraphs):
    if '•' in para.text and 'Bullet point example' in para.text:
        left_indent = para.paragraph_format.left_indent
        print(f"Bullet: '{para.text}'")
        print(f"  Left indent: {left_indent}")
        break

# 3. Check for sample image
print("\n3. SAMPLE IMAGE IN IMAGE SECTION:")
print("-" * 80)

image_found = False
for idx, para in enumerate(doc.paragraphs):
    if 'Image Section' in para.text:
        # Check next 10 paragraphs for image
        for i in range(idx, min(idx + 10, len(doc.paragraphs))):
            check_para = doc.paragraphs[i]
            for run in check_para.runs:
                if hasattr(run, '_element'):
                    drawings = run._element.findall('.//{http://schemas.openxmlformats.org/wordprocessingml/2006/main}drawing')
                    blips = run._element.findall('.//{http://schemas.openxmlformats.org/drawingml/2006/main}blip')
                    if drawings or blips:
                        image_found = True
                        # Get image dimensions
                        extents = run._element.findall('.//{http://schemas.openxmlformats.org/drawingml/2006/wordprocessingDrawing}extent')
                        if extents:
                            for extent in extents:
                                cx = extent.get('cx')
                                cy = extent.get('cy')
                                if cx and cy:
                                    width_inches = int(cx) / 914400
                                    height_inches = int(cy) / 914400
                                    print(f"✅ Image found in Section 3")
                                    print(f"   Dimensions: {width_inches:.2f}\" × {height_inches:.2f}\"")
                        break
            if image_found:
                break
        break

if not image_found:
    print("❌ No image found in Image Section")

# 4. Check List of Figures
print("\n4. LIST OF FIGURES:")
print("-" * 80)

lof_found = False
for idx, para in enumerate(doc.paragraphs):
    if 'List of Figures' in para.text and para.style.name == 'Heading 1':
        lof_found = True
        # Check next paragraphs for figure entries
        for i in range(idx + 1, min(idx + 5, len(doc.paragraphs))):
            lof_para = doc.paragraphs[i]
            if 'Figure' in lof_para.text:
                print(f"Figure entry: '{lof_para.text}'")
                break
        break

if not lof_found:
    print("❌ List of Figures not found")

print("\n" + "=" * 80)
print("VERIFICATION COMPLETE")
print("=" * 80)

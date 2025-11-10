#!/usr/bin/env python3
"""
Create Word template that matches the working document formatting.
Based on reverse-engineering: aws/cloud/onpremise-to-cloud-migration/presales/statement-of-work.docx
"""

import argparse
from docx import Document
from docx.shared import Pt, RGBColor, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_BREAK
from docx.enum.table import WD_ALIGN_VERTICAL
from docx.oxml.shared import OxmlElement
from docx.oxml.ns import qn
from pathlib import Path


def configure_styles(doc):
    """Configure document styles to match reference document."""

    # Set document defaults to 12pt Times New Roman
    # This ensures all text defaults to 12pt unless specifically overridden
    styles_part = None
    for rel in doc.part.rels.values():
        if 'styles' in rel.target_ref:
            styles_part = rel.target_part
            break

    if styles_part:
        styles_element = styles_part.element
        doc_defaults = styles_element.find(qn('w:docDefaults'))

        if doc_defaults is None:
            doc_defaults = OxmlElement('w:docDefaults')
            styles_element.insert(0, doc_defaults)

        # Set rPrDefault (run property defaults)
        rpr_default = doc_defaults.find(qn('w:rPrDefault'))
        if rpr_default is None:
            rpr_default = OxmlElement('w:rPrDefault')
            doc_defaults.append(rpr_default)

        rpr = rpr_default.find(qn('w:rPr'))
        if rpr is None:
            rpr = OxmlElement('w:rPr')
            rpr_default.append(rpr)

        # Set font to Times New Roman
        rfonts = rpr.find(qn('w:rFonts'))
        if rfonts is None:
            rfonts = OxmlElement('w:rFonts')
            rpr.append(rfonts)
        rfonts.set(qn('w:ascii'), 'Times New Roman')
        rfonts.set(qn('w:hAnsi'), 'Times New Roman')
        rfonts.set(qn('w:cs'), 'Times New Roman')

        # Set default font size to 12pt (24 half-points)
        sz = rpr.find(qn('w:sz'))
        if sz is None:
            sz = OxmlElement('w:sz')
            rpr.append(sz)
        sz.set(qn('w:val'), '24')  # 24 half-points = 12pt

        # Set default font size for complex scripts
        sz_cs = rpr.find(qn('w:szCs'))
        if sz_cs is None:
            sz_cs = OxmlElement('w:szCs')
            rpr.append(sz_cs)
        sz_cs.set(qn('w:val'), '24')  # 24 half-points = 12pt

    # Normal style - Times New Roman 12pt (default body text)
    normal_style = doc.styles['Normal']
    normal_font = normal_style.font
    normal_font.name = 'Times New Roman'
    normal_font.size = Pt(12)

    # Title style - 28pt Bold #1F4E78
    title_style = doc.styles['Title']
    title_font = title_style.font
    title_font.name = 'Times New Roman'
    title_font.size = Pt(28)
    title_font.bold = True
    title_font.color.rgb = RGBColor(31, 78, 120)  # #1F4E78

    # Heading 1 - 16pt Bold #1F4E78 SMALL CAPS
    h1_style = doc.styles['Heading 1']
    h1_font = h1_style.font
    h1_font.name = 'Times New Roman'
    h1_font.size = Pt(16)
    h1_font.bold = True
    h1_font.color.rgb = RGBColor(31, 78, 120)  # #1F4E78
    h1_font.small_caps = True

    # Heading 2 - Bold #2E5C8A SMALL CAPS
    h2_style = doc.styles['Heading 2']
    h2_font = h2_style.font
    h2_font.name = 'Times New Roman'
    h2_font.bold = True
    h2_font.color.rgb = RGBColor(46, 92, 138)  # #2E5C8A
    h2_font.small_caps = True

    # Heading 3 - Bold #404040 SMALL CAPS
    h3_style = doc.styles['Heading 3']
    h3_font = h3_style.font
    h3_font.name = 'Times New Roman'
    h3_font.bold = True
    h3_font.color.rgb = RGBColor(64, 64, 64)  # #404040
    h3_font.small_caps = True


def set_cell_border(cell, **kwargs):
    """
    Set cell border.
    Usage:
        set_cell_border(
            cell,
            top={"sz": 1, "val": "single", "color": "CCCCCC"},
            bottom={"sz": 1, "val": "single", "color": "CCCCCC"},
            left={"sz": 1, "val": "single", "color": "CCCCCC"},
            right={"sz": 1, "val": "single", "color": "CCCCCC"},
        )
    """
    tc = cell._element
    tcPr = tc.get_or_add_tcPr()

    # Create border elements
    tcBorders = OxmlElement('w:tcBorders')

    for edge in ('left', 'top', 'right', 'bottom'):
        if edge in kwargs:
            edge_data = kwargs.get(edge)
            edge_el = OxmlElement(f'w:{edge}')
            for key, value in edge_data.items():
                edge_el.set(qn(f'w:{key}'), str(value))
            tcBorders.append(edge_el)

    tcPr.append(tcBorders)

    # Add cell margins for padding (top/bottom/left/right)
    tcMar = OxmlElement('w:tcMar')
    for margin_side in ['top', 'bottom', 'left', 'right']:
        margin_elem = OxmlElement(f'w:{margin_side}')
        margin_elem.set(qn('w:w'), '80')  # 80 twips = approx 0.06 inches (moderate padding)
        margin_elem.set(qn('w:type'), 'dxa')
        tcMar.append(margin_elem)
    tcPr.append(tcMar)


def set_cell_background(cell, color):
    """Set cell background color."""
    shading_elm = OxmlElement('w:shd')
    shading_elm.set(qn('w:fill'), color)
    cell._element.get_or_add_tcPr().append(shading_elm)


def create_metadata_table(doc):
    """Create 2-column metadata table for cover page with borders and formatting."""
    table = doc.add_table(rows=5, cols=2)
    table.style = 'Normal Table'

    # Set column widths to match reference document exactly
    # Reference: Column 1 = 2.094", Column 2 = 4.172"
    table.columns[0].width = Inches(2.094)
    table.columns[1].width = Inches(4.172)

    # Define metadata fields
    fields = [
        ('Client Name:', '[Client Company Name]'),
        ('Date:', '[MM/DD/YYYY]'),
        ('Consulting Company:', '[Consulting Company Name]'),
        ('Submitted By:', '[Full Name]\n[email@example.com]\n[Phone: XXX-XXX-XXXX]'),
        ('Version:', 'v1.0')
    ]

    # Border definition - light gray 1pt borders on all sides
    border_spec = {
        "sz": "1",
        "val": "single",
        "color": "CCCCCC",
        "space": "0"
    }

    # Populate table
    for row_idx, (label, value) in enumerate(fields):
        row = table.rows[row_idx]

        # Label cell (bold, very light blue background)
        # Note: Reference doc does NOT have vertical alignment set
        label_cell = row.cells[0]
        label_cell.width = Inches(2.094)  # Set cell width
        label_para = label_cell.paragraphs[0]
        label_run = label_para.add_run(label)
        label_run.font.name = 'Times New Roman'
        label_run.font.bold = True
        label_run.font.size = Pt(12)
        # Set line spacing to 1.0 and add top/bottom padding
        label_para.paragraph_format.line_spacing = 1.0
        label_para.paragraph_format.space_before = Pt(6)
        label_para.paragraph_format.space_after = Pt(6)

        # Set very light blue background on label cell (E8F0F8 - very light blue)
        set_cell_background(label_cell, 'E8F0F8')

        # Set borders on label cell
        set_cell_border(
            label_cell,
            top=border_spec,
            bottom=border_spec,
            left=border_spec,
            right=border_spec
        )

        # Value cell
        # Note: Reference doc does NOT have vertical alignment set
        value_cell = row.cells[1]
        value_cell.width = Inches(4.172)  # Set cell width
        value_para = value_cell.paragraphs[0]
        value_run = value_para.add_run(value)
        value_run.font.name = 'Times New Roman'
        value_run.font.size = Pt(12)
        # Set line spacing to 1.0 and add top/bottom padding
        value_para.paragraph_format.line_spacing = 1.0
        value_para.paragraph_format.space_before = Pt(6)
        value_para.paragraph_format.space_after = Pt(6)

        # Set borders on value cell
        set_cell_border(
            value_cell,
            top=border_spec,
            bottom=border_spec,
            left=border_spec,
            right=border_spec
        )

    return table


def add_page_break(paragraph):
    """Add a page break to a paragraph."""
    run = paragraph.add_run()
    run.add_break(break_type=WD_BREAK.PAGE)


def add_page_number(paragraph):
    """Add page number field to paragraph."""
    run = paragraph.add_run()
    fldChar1 = OxmlElement('w:fldChar')
    fldChar1.set(qn('w:fldCharType'), 'begin')

    instrText = OxmlElement('w:instrText')
    instrText.set(qn('xml:space'), 'preserve')
    instrText.text = 'PAGE'

    fldChar2 = OxmlElement('w:fldChar')
    fldChar2.set(qn('w:fldCharType'), 'end')

    run._r.append(fldChar1)
    run._r.append(instrText)
    run._r.append(fldChar2)
    return run


def add_footer(doc, logo_path):
    """Add footer with document name, page number, and EO Framework logo."""
    section = doc.sections[0]
    footer = section.footer

    # Clear existing footer paragraphs
    footer_para = footer.paragraphs[0] if footer.paragraphs else footer.add_paragraph()
    footer_para.alignment = WD_ALIGN_PARAGRAPH.LEFT

    # Add document name on left
    doc_name_run = footer_para.add_run('[Document Name]')
    doc_name_run.font.name = 'Times New Roman'
    doc_name_run.font.size = Pt(12)  # Changed from 9pt to 12pt
    doc_name_run.font.color.rgb = RGBColor(128, 128, 128)

    # Add tab to center
    footer_para.add_run('\t')

    # Add page number in center
    page_label = footer_para.add_run('Page ')
    page_label.font.name = 'Times New Roman'
    page_label.font.size = Pt(12)  # Changed from 9pt to 12pt
    page_label.font.color.rgb = RGBColor(128, 128, 128)

    add_page_number(footer_para)

    # Add tab to right
    footer_para.add_run('\t')

    # Add EO Framework logo on right
    eo_logo = logo_path / 'eo-framework-logo-real.png'
    if eo_logo.exists():
        logo_run = footer_para.add_run()
        logo_run.add_picture(str(eo_logo), height=Pt(12))

    # Set tab stops: center at 3.25", right at 6.5"
    from docx.enum.text import WD_TAB_ALIGNMENT
    tab_stops = footer_para.paragraph_format.tab_stops
    tab_stops.add_tab_stop(Inches(3.25), WD_TAB_ALIGNMENT.CENTER)
    tab_stops.add_tab_stop(Inches(6.5), WD_TAB_ALIGNMENT.RIGHT)

    # Set paragraph spacing
    footer_para.paragraph_format.space_before = Pt(0)
    footer_para.paragraph_format.space_after = Pt(0)
    footer_para.paragraph_format.line_spacing = 1.0


def create_word_template(template_type='general', output_name=None):
    """
    Create Word template matching the reference document structure.

    Args:
        template_type: 'general' or 'sow' (currently both use same structure)
        output_name: Custom output filename (default: EOFramework-Word-Template-01.docx)
    """
    if output_name:
        output_file = Path(f'support/doc-templates/word/{output_name}')
    else:
        output_file = Path('support/doc-templates/word/EOFramework-Word-Template-01.docx')
    logo_path = Path('support/doc-templates/assets/logos')

    print("=" * 80)
    print("CREATING WORD TEMPLATE (Reference-based)")
    print("=" * 80)

    # Create new document
    doc = Document()

    # Configure styles first
    print("\n📐 Configuring styles...")
    configure_styles(doc)
    print("  ✅ Normal: Times New Roman 12pt")
    print("  ✅ Title: Times New Roman 28pt Bold #1F4E78")
    print("  ✅ Heading 1: Times New Roman 16pt Bold #1F4E78 SMALL CAPS")
    print("  ✅ Heading 2: Times New Roman Bold #2E5C8A SMALL CAPS")
    print("  ✅ Heading 3: Times New Roman Bold #404040 SMALL CAPS")

    # =========================================================================
    # COVER PAGE
    # =========================================================================
    print("\n📋 Creating cover page...")

    # Spacing at top
    doc.add_paragraph()

    # Logo (centered, 4.17" x 1.04")
    logo_para = doc.add_paragraph()
    logo_para.alignment = WD_ALIGN_PARAGRAPH.CENTER

    consulting_logo = logo_path / 'consulting_company_logo.png'
    if consulting_logo.exists():
        logo_run = logo_para.add_run()
        logo_run.add_picture(str(consulting_logo), width=Inches(4.17))
        print("  ✅ Consulting company logo added (4.17\" x 1.04\")")
    else:
        print("  ⚠️  Logo not found: consulting_company_logo.png")

    # Spacing
    doc.add_paragraph()

    # Title (centered)
    title = doc.add_paragraph('Statement Of Work', style='Title')
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    print("  ✅ Title added (centered)")

    # Subtitle (centered, normal style)
    subtitle = doc.add_paragraph()
    subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
    subtitle_run = subtitle.add_run('[Project Name] Statement of Work (SOW)')
    subtitle_run.font.name = 'Times New Roman'
    subtitle_run.font.size = Pt(12)
    print("  ✅ Subtitle added")

    # Spacing
    doc.add_paragraph()
    doc.add_paragraph()

    # Metadata table with borders and light blue/grey background on labels
    create_metadata_table(doc)
    print("  ✅ Metadata table created with:")
    print("      - Gray borders (#CCCCCC)")
    print("      - Very light blue background on labels (#E8F0F8)")
    print("      - No vertical alignment (default)")
    print("      - Column widths: 2.094\" / 4.172\"")
    print("      - 12pt font, line spacing 1.0, 6pt padding")

    # Spacing after table
    centered_space = doc.add_paragraph(' ')
    centered_space.alignment = WD_ALIGN_PARAGRAPH.CENTER

    # Empty paragraph before page break
    empty_para = doc.add_paragraph()

    # PAGE BREAK before Table of Contents
    add_page_break(empty_para)
    print("  ✅ Page break added before TOC")

    # =========================================================================
    # TABLE OF CONTENTS (NEW PAGE)
    # =========================================================================
    print("\n📑 Creating Table of Contents section (new page)...")

    # Empty paragraph after page break
    doc.add_paragraph()

    toc_heading = doc.add_paragraph('Table of Contents', style='Heading 1')
    doc.add_paragraph()

    # Insert dynamic TOC field
    # This creates a real TOC that Word can update
    toc_para = doc.add_paragraph()
    toc_run = toc_para.add_run()

    # Create the TOC field XML
    # This tells Word to create a TOC with headings 1-3, showing page numbers with dot leaders
    fldChar_begin = OxmlElement('w:fldChar')
    fldChar_begin.set(qn('w:fldCharType'), 'begin')

    instrText = OxmlElement('w:instrText')
    instrText.set(qn('xml:space'), 'preserve')
    instrText.text = r'TOC \o "1-3" \h \z \u'  # TOC with outline levels 1-3

    fldChar_separate = OxmlElement('w:fldChar')
    fldChar_separate.set(qn('w:fldCharType'), 'separate')

    fldChar_end = OxmlElement('w:fldChar')
    fldChar_end.set(qn('w:fldCharType'), 'end')

    # Add field begin, instruction, and separate to the first run
    r_element = toc_run._element
    r_element.append(fldChar_begin)
    r_element.append(instrText)
    r_element.append(fldChar_separate)

    # Add initial TOC entries as runs within the same paragraph (between separate and end markers)
    # This creates field result content that users can right-click to update
    toc_entries = [
        "1 Introduction",
        "    1.1 Purpose",
        "        1.1.1 Scope",
        "2 Table Content",
        "3 Image Section",
        "4 Additional Content",
        "5 Sign-Off"
    ]

    # Add TOC content as runs with line breaks in the field paragraph
    for idx, entry in enumerate(toc_entries):
        if idx > 0:
            # Add line break before each entry (except first)
            toc_para.add_run().add_break()
        entry_run = toc_para.add_run(entry)
        entry_run.font.name = 'Times New Roman'
        entry_run.font.size = Pt(12)

    # Add the end field marker in a new run at the end
    end_run = toc_para.add_run()
    end_run._element.append(fldChar_end)

    # Format TOC paragraph (the field container)
    toc_para.paragraph_format.line_spacing = 1.15
    toc_para.paragraph_format.space_before = Pt(3)
    toc_para.paragraph_format.space_after = Pt(3)

    print("  ✅ TOC heading added")
    print("  ✅ Dynamic TOC field inserted (will auto-generate from headings)")

    # Page break before List of Figures
    empty_para2 = doc.add_paragraph()
    add_page_break(empty_para2)

    # =========================================================================
    # LIST OF FIGURES (NEW PAGE)
    # =========================================================================
    print("\n📊 Creating List of Figures section (new page)...")

    doc.add_paragraph()
    lof_heading = doc.add_paragraph('List of Figures', style='Heading 1')
    doc.add_paragraph()

    # Insert dynamic List of Figures field (TOC for figures only)
    lof_para = doc.add_paragraph()
    lof_run = lof_para.add_run()

    # Create TOC field for figures (captions with "Figure" label)
    fldChar_begin_lof = OxmlElement('w:fldChar')
    fldChar_begin_lof.set(qn('w:fldCharType'), 'begin')

    instrText_lof = OxmlElement('w:instrText')
    instrText_lof.set(qn('xml:space'), 'preserve')
    instrText_lof.text = r'TOC \h \z \c "Figure"'  # TOC for items with "Figure" caption label

    fldChar_separate_lof = OxmlElement('w:fldChar')
    fldChar_separate_lof.set(qn('w:fldCharType'), 'separate')

    fldChar_end_lof = OxmlElement('w:fldChar')
    fldChar_end_lof.set(qn('w:fldCharType'), 'end')

    # Add field begin, instruction, and separate to the first run
    r_element_lof = lof_run._element
    r_element_lof.append(fldChar_begin_lof)
    r_element_lof.append(instrText_lof)
    r_element_lof.append(fldChar_separate_lof)

    # Add initial figure entry as run within the same paragraph (between separate and end markers)
    figure_run = lof_para.add_run("Figure 1: Sample Architecture Diagram")
    figure_run.font.name = 'Times New Roman'
    figure_run.font.size = Pt(12)

    # Add the end field marker in a new run at the end
    end_run_lof = lof_para.add_run()
    end_run_lof._element.append(fldChar_end_lof)

    # Format List of Figures paragraph (the field container)
    lof_para.paragraph_format.line_spacing = 1.0
    lof_para.paragraph_format.space_before = Pt(3)
    lof_para.paragraph_format.space_after = Pt(3)

    print("  ✅ List of Figures heading added")
    print("  ✅ Dynamic List of Figures field inserted (will auto-generate from figure captions)")

    # Page break before content sections
    empty_para3 = doc.add_paragraph()
    add_page_break(empty_para3)

    # =========================================================================
    # CONTENT SECTIONS (NEW PAGE)
    # =========================================================================
    print("\n📝 Creating content sections (new page)...")

    # Section 1 - with multi-level sub-headers
    doc.add_paragraph('1 Section Header 1', style='Heading 1')
    para1 = doc.add_paragraph()
    para1_run = para1.add_run('This is the content for Section Header 1.')
    para1_run.font.name = 'Times New Roman'
    para1_run.font.size = Pt(12)

    # 1.1
    doc.add_paragraph('1.1 Section Sub-Header 1', style='Heading 2')
    para11 = doc.add_paragraph()
    para11_run = para11.add_run('Content for Section Sub-Header 1.')
    para11_run.font.name = 'Times New Roman'
    para11_run.font.size = Pt(12)

    # 1.1.1
    doc.add_paragraph('1.1.1 Section Sub-Header 2', style='Heading 3')
    para111 = doc.add_paragraph()
    para111_run = para111.add_run('Content for Section Sub-Header 2.')
    para111_run.font.name = 'Times New Roman'
    para111_run.font.size = Pt(12)

    # Bullet list example with left indent and reduced line spacing
    bullet1 = doc.add_paragraph()
    bullet_run1 = bullet1.add_run('• Bullet point example 1')
    bullet_run1.font.name = 'Times New Roman'
    bullet_run1.font.size = Pt(12)
    bullet1.paragraph_format.left_indent = Inches(0.25)
    bullet1.paragraph_format.line_spacing = 1.0
    bullet1.paragraph_format.space_after = Pt(0)

    bullet2 = doc.add_paragraph()
    bullet_run2 = bullet2.add_run('• Bullet point example 2')
    bullet_run2.font.name = 'Times New Roman'
    bullet_run2.font.size = Pt(12)
    bullet2.paragraph_format.left_indent = Inches(0.25)
    bullet2.paragraph_format.line_spacing = 1.0
    bullet2.paragraph_format.space_after = Pt(0)

    # 1.1.1.1 (using Heading 3 style - Word doesn't have default Heading 4+)
    doc.add_paragraph('1.1.1.1 Section Sub-Header 3', style='Heading 3')
    para1111 = doc.add_paragraph()
    para1111_run = para1111.add_run('Content for Section Sub-Header 3.')
    para1111_run.font.name = 'Times New Roman'
    para1111_run.font.size = Pt(12)

    # 1.1.1.2
    doc.add_paragraph('1.1.1.2 Section Sub-Header 4', style='Heading 3')
    para1112 = doc.add_paragraph()
    para1112_run = para1112.add_run('Content for Section Sub-Header 4.')
    para1112_run.font.name = 'Times New Roman'
    para1112_run.font.size = Pt(12)

    # 1.2
    doc.add_paragraph('1.2 Section Sub-Header 5', style='Heading 2')
    para12 = doc.add_paragraph()
    para12_run = para12.add_run('Content for Section Sub-Header 5.')
    para12_run.font.name = 'Times New Roman'
    para12_run.font.size = Pt(12)

    print("  ✅ Section 1 with multi-level headers added")

    # =========================================================================
    # SECTION 2 - TABLE CONTENT
    # =========================================================================
    print("\n📊 Creating Section 2 - Table Content...")

    doc.add_paragraph('2 Section Header 2 - Table Content', style='Heading 1')

    # 2.1
    doc.add_paragraph('2.1 Table Section 1', style='Heading 2')
    para21 = doc.add_paragraph()
    para21_run = para21.add_run('Example table section with sample deliverables table.')
    para21_run.font.name = 'Times New Roman'
    para21_run.font.size = Pt(12)

    # Add spacing before table
    spacing_para = doc.add_paragraph()
    spacing_para.paragraph_format.space_before = Pt(6)
    spacing_para.paragraph_format.space_after = Pt(0)

    # Create sample table
    table1 = doc.add_table(rows=4, cols=5)
    table1.style = 'Normal Table'

    # Header row with light blue background
    header_cells = table1.rows[0].cells
    headers = ['#', 'Deliverable', 'Type', 'Due Date', 'Acceptance']
    for idx, header_text in enumerate(headers):
        cell = header_cells[idx]
        # Set vertical alignment to center
        cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
        para = cell.paragraphs[0]
        # Add padding to header cells
        para.paragraph_format.space_before = Pt(6)
        para.paragraph_format.space_after = Pt(6)
        run = para.add_run(header_text)
        run.font.name = 'Times New Roman'
        run.font.bold = True
        # Set light blue background matching metadata table
        set_cell_background(cell, 'E8F0F8')
        # Add borders
        border_spec = {"sz": "1", "val": "single", "color": "CCCCCC", "space": "0"}
        set_cell_border(cell, top=border_spec, bottom=border_spec, left=border_spec, right=border_spec)

    # Sample data rows
    sample_data = [
        ['1', 'Assessment Report', 'Document', '[Date]', '[Client Lead]'],
        ['2', 'Architecture Design', 'Document', '[Date]', '[Architect]'],
        ['3', 'Implementation Guide', 'Document', '[Date]', '[Team Lead]']
    ]

    for row_idx, row_data in enumerate(sample_data, start=1):
        cells = table1.rows[row_idx].cells
        for col_idx, cell_text in enumerate(row_data):
            cell = cells[col_idx]
            para = cell.paragraphs[0]
            # Add equal top and bottom padding to data cells for symmetry
            para.paragraph_format.space_before = Pt(4)
            para.paragraph_format.space_after = Pt(4)
            run = para.add_run(cell_text)
            run.font.name = 'Times New Roman'
            # Add borders to content cells
            border_spec = {"sz": "1", "val": "single", "color": "CCCCCC", "space": "0"}
            set_cell_border(cell, top=border_spec, bottom=border_spec, left=border_spec, right=border_spec)

    doc.add_paragraph()  # Spacing after table

    # 2.2
    doc.add_paragraph('2.2 Table Section 2', style='Heading 2')
    para22 = doc.add_paragraph()
    para22_run = para22.add_run('Another table example - RACI matrix or similar.')
    para22_run.font.name = 'Times New Roman'
    para22_run.font.size = Pt(12)

    # Add spacing before table
    spacing_para2 = doc.add_paragraph()
    spacing_para2.paragraph_format.space_before = Pt(6)
    spacing_para2.paragraph_format.space_after = Pt(0)

    # Create RACI-style table
    table2 = doc.add_table(rows=4, cols=4)
    table2.style = 'Normal Table'

    # Header row with light blue background
    raci_headers = ['Task/Activity', 'Team A', 'Team B', 'Team C']
    for idx, header_text in enumerate(raci_headers):
        cell = table2.rows[0].cells[idx]
        # Set vertical alignment to center
        cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
        para = cell.paragraphs[0]
        # Add padding to header cells
        para.paragraph_format.space_before = Pt(6)
        para.paragraph_format.space_after = Pt(6)
        run = para.add_run(header_text)
        run.font.name = 'Times New Roman'
        run.font.bold = True
        # Set light blue background
        set_cell_background(cell, 'E8F0F8')
        # Add borders
        border_spec = {"sz": "1", "val": "single", "color": "CCCCCC", "space": "0"}
        set_cell_border(cell, top=border_spec, bottom=border_spec, left=border_spec, right=border_spec)

    # Sample RACI data
    raci_data = [
        ['Activity 1', 'R', 'A', 'C'],
        ['Activity 2', 'A', 'R', 'I'],
        ['Activity 3', 'C', 'I', 'R']
    ]

    for row_idx, row_data in enumerate(raci_data, start=1):
        for col_idx, cell_text in enumerate(row_data):
            cell = table2.rows[row_idx].cells[col_idx]
            para = cell.paragraphs[0]
            # Add equal top and bottom padding to data cells for symmetry
            para.paragraph_format.space_before = Pt(4)
            para.paragraph_format.space_after = Pt(4)
            run = para.add_run(cell_text)
            run.font.name = 'Times New Roman'
            # Add borders to content cells
            border_spec = {"sz": "1", "val": "single", "color": "CCCCCC", "space": "0"}
            set_cell_border(cell, top=border_spec, bottom=border_spec, left=border_spec, right=border_spec)

    print("  ✅ Section 2 with table examples added")

    # =========================================================================
    # SECTION 3 - IMAGE SECTION
    # =========================================================================
    print("\n🖼️  Creating Section 3 - Image Section...")

    doc.add_paragraph()  # Spacing
    doc.add_paragraph('3 Section Header 3 - Image Section', style='Heading 1')

    # 3.1
    doc.add_paragraph('3.1 Image Content', style='Heading 2')
    para31 = doc.add_paragraph()
    para31_run = para31.add_run('Below is a sample architecture diagram showing the relationship between system components.')
    para31_run.font.name = 'Times New Roman'
    para31_run.font.size = Pt(12)

    # Insert sample diagram
    sample_diagram_path = logo_path.parent / 'sample-diagram.png'
    if sample_diagram_path.exists():
        img_para = doc.add_paragraph()
        img_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
        img_run = img_para.add_run()
        img_run.add_picture(str(sample_diagram_path), width=Inches(5.0))

        # Add figure caption
        caption_para = doc.add_paragraph()
        caption_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
        caption_run = caption_para.add_run('Figure 1: Sample Architecture Diagram')
        caption_run.font.name = 'Times New Roman'
        caption_run.font.size = Pt(12)  # Changed from 11pt to 12pt
        caption_run.font.italic = True
        caption_run.font.color.rgb = RGBColor(64, 64, 64)
        print("  ✅ Section 3 with sample diagram added")
    else:
        # Fallback placeholder
        img_placeholder = doc.add_paragraph()
        img_placeholder.alignment = WD_ALIGN_PARAGRAPH.CENTER
        img_run = img_placeholder.add_run('[DIAGRAM OR IMAGE PLACEHOLDER]')
        img_run.font.name = 'Times New Roman'
        img_run.font.size = Pt(12)
        img_run.font.italic = True
        img_run.font.color.rgb = RGBColor(128, 128, 128)
        print("  ⚠️  Section 3 with placeholder (sample-diagram.png not found)")

    # =========================================================================
    # SECTION 4 - ADDITIONAL CONTENT
    # =========================================================================
    print("\n📝 Creating Section 4 - Additional Content...")

    doc.add_paragraph()  # Spacing
    doc.add_paragraph('4 Section Header 4', style='Heading 1')
    para4 = doc.add_paragraph()
    para4_run = para4.add_run('Additional section content.')
    para4_run.font.name = 'Times New Roman'
    para4_run.font.size = Pt(12)

    # 4.1
    doc.add_paragraph('4.1 Additional Sub-Section', style='Heading 2')
    para41 = doc.add_paragraph()
    para41_run = para41.add_run('Content for additional sub-section.')
    para41_run.font.name = 'Times New Roman'
    para41_run.font.size = Pt(12)

    print("  ✅ Section 4 added")

    # =========================================================================
    # SECTION 5 - APPROVALS AND SIGNATURES
    # =========================================================================
    print("\n✍️  Creating Section 5 - Sign-Off...")

    doc.add_paragraph()  # Spacing
    doc.add_paragraph('5 Sign-Off', style='Heading 1')
    doc.add_paragraph()  # Spacing

    # Signature table with proper formatting and padding
    sig_table = doc.add_table(rows=5, cols=2)
    sig_table.style = 'Normal Table'

    # Define border for signature table
    border_spec = {"sz": "1", "val": "single", "color": "CCCCCC", "space": "0"}

    # Row labels and structure
    row_labels = ['Party', 'Signature', 'Name', 'Title', 'Date']

    for row_idx, label in enumerate(row_labels):
        row = sig_table.rows[row_idx]

        # Left column - Client
        left_cell = row.cells[0]
        left_para = left_cell.paragraphs[0]

        if row_idx == 0:  # Header row
            left_para.alignment = WD_ALIGN_PARAGRAPH.CENTER  # Center align content
            run = left_para.add_run('CLIENT')
            run.font.bold = True
            run.font.name = 'Times New Roman'
            run.font.size = Pt(12)
            set_cell_background(left_cell, 'E8F0F8')
            left_para.paragraph_format.space_before = Pt(6)
            left_para.paragraph_format.space_after = Pt(6)
        else:
            # All signature fields - left aligned, bottom justified, same padding
            left_para.alignment = WD_ALIGN_PARAGRAPH.LEFT
            run = left_para.add_run(f'{label}: ____________________')
            left_para.paragraph_format.space_before = Pt(18)
            left_para.paragraph_format.space_after = Pt(6)
            # Set vertical alignment to bottom
            left_cell.vertical_alignment = WD_ALIGN_VERTICAL.BOTTOM
            left_para.runs[0].font.name = 'Times New Roman'
            left_para.runs[0].font.size = Pt(12)

        # Set borders
        set_cell_border(left_cell, top=border_spec, bottom=border_spec, left=border_spec, right=border_spec)

        # Right column - Consulting Company
        right_cell = row.cells[1]
        right_para = right_cell.paragraphs[0]

        if row_idx == 0:  # Header row
            right_para.alignment = WD_ALIGN_PARAGRAPH.CENTER  # Center align content
            run = right_para.add_run('CONSULTING COMPANY')
            run.font.bold = True
            run.font.name = 'Times New Roman'
            run.font.size = Pt(12)
            set_cell_background(right_cell, 'E8F0F8')
            right_para.paragraph_format.space_before = Pt(6)
            right_para.paragraph_format.space_after = Pt(6)
        else:
            # All signature fields - left aligned, bottom justified, same padding
            right_para.alignment = WD_ALIGN_PARAGRAPH.LEFT
            run = right_para.add_run(f'{label}: ____________________')
            right_para.paragraph_format.space_before = Pt(18)
            right_para.paragraph_format.space_after = Pt(6)
            # Set vertical alignment to bottom
            right_cell.vertical_alignment = WD_ALIGN_VERTICAL.BOTTOM
            right_para.runs[0].font.name = 'Times New Roman'
            right_para.runs[0].font.size = Pt(12)

        # Set borders
        set_cell_border(right_cell, top=border_spec, bottom=border_spec, left=border_spec, right=border_spec)

    # Set row heights for signature row (more space)
    sig_row = sig_table.rows[1]  # Signature row
    tr = sig_row._element
    trPr = tr.get_or_add_trPr()
    trHeight = OxmlElement('w:trHeight')
    trHeight.set(qn('w:val'), '720')  # 0.5 inch in twentieths of a point
    trHeight.set(qn('w:hRule'), 'atLeast')
    trPr.append(trHeight)

    # Add spacing after signature table for visual separation
    spacing_para = doc.add_paragraph()
    spacing_para.paragraph_format.space_after = Pt(12)

    print("  ✅ Section 5 with formatted signature table added (centered content, 24pt signature padding)")

    # =========================================================================
    # FOOTER: Document name, page number, and EO Framework logo
    # =========================================================================
    print("\n📄 Adding footer...")

    add_footer(doc, logo_path)
    print("  ✅ Footer added with:")
    print("      - Document name (left)")
    print("      - Page number (center)")
    print("      - EO Framework logo (right)")

    # =========================================================================
    # SAVE TEMPLATE
    # =========================================================================
    print("\n" + "=" * 80)
    print("💾 Saving template...")

    output_file.parent.mkdir(parents=True, exist_ok=True)
    doc.save(str(output_file))

    print(f"✅ Template saved: {output_file}")
    print("\n" + "=" * 80)
    print("✨ TEMPLATE CREATION COMPLETE")
    print("=" * 80)
    print("\nTemplate structure:")
    print("  📋 Page 1: Cover with logo and metadata table")
    print("  📑 Page 2: Table of Contents (with all section entries)")
    print("  📊 Page 3: List of Figures")
    print("  📝 Page 4+: Content sections:")
    print("      • Section 1: Multi-level headers (1, 1.1, 1.1.1, 1.1.1.1, etc.)")
    print("      • Section 2: Table content examples")
    print("      • Section 3: Image section placeholder")
    print("      • Section 4: Additional content")
    print("      • Section 5: Sign-Off")
    print("\nFormatting:")
    print("  • All fonts: Times New Roman")
    print("  • Title: 28pt Bold #1F4E78 (centered)")
    print("  • Heading 1: 16pt Bold #1F4E78 SMALL CAPS")
    print("  • Heading 2: Bold #2E5C8A SMALL CAPS")
    print("  • Heading 3: Bold #404040 SMALL CAPS")
    print("  • Body text: 12pt")
    print("  • Metadata table:")
    print("      - Gray borders (#CCCCCC)")
    print("      - Very light blue background on labels (#E8F0F8)")
    print("      - No vertical alignment (default)")
    print("      - Column widths: 2.094\" / 4.172\"")
    print("      - 12pt font, line spacing 1.0, 6pt padding")
    print("  • Table of Contents: 12pt font, line spacing 1.15, 3pt spacing")
    print("  • List of Figures: Sample architecture diagram entry")
    print("  • Bullet points: 0.25\" left indent, line spacing 1.0, no space after")
    print("  • Content tables: Light blue headers (#E8F0F8), vertically centered, gray borders")
    print("  • Signature table: Centered content, 24pt signature padding, 12pt other rows")
    print("  • Footer: Document name | Page # | EO Framework logo (all 12pt)")
    print()


# ============================================================================
# CLI INTERFACE
# ============================================================================

if __name__ == '__main__':
    parser = argparse.ArgumentParser(
        description='Create Word template matching reference document formatting.'
    )
    parser.add_argument(
        '--type',
        choices=['general', 'sow'],
        default='general',
        help='Template type (currently both use same structure)'
    )
    parser.add_argument(
        '--output',
        type=str,
        help='Custom output filename (e.g., My-Template.docx)'
    )

    args = parser.parse_args()
    create_word_template(args.type, args.output)

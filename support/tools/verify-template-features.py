#!/usr/bin/env python3
"""
Verify specific template features.
"""

from docx import Document
from pathlib import Path
import sys

docx_path = Path(sys.argv[1])
doc = Document(docx_path)

print("=" * 80)
print("TEMPLATE FEATURE VERIFICATION")
print("=" * 80)

# 1. Check heading styles for small caps
print("\n1. HEADING STYLES - SMALL CAPS:")
print("-" * 80)
for style_name in ['Heading 1', 'Heading 2', 'Heading 3']:
    style = doc.styles[style_name]
    font = style.font
    print(f"{style_name}:")
    print(f"  Small caps: {font.small_caps}")
    print(f"  Bold: {font.bold}")
    print(f"  Color: {font.color.rgb if font.color and font.color.rgb else 'Not set'}")

# 2. Check content tables for header background
print("\n2. CONTENT TABLES - HEADER BACKGROUND:")
print("-" * 80)

table_count = 0
for idx, table in enumerate(doc.tables):
    # Skip metadata table (first table)
    if idx == 0:
        continue

    table_count += 1
    print(f"\nTable {table_count} ({len(table.rows)} rows × {len(table.columns)} cols):")

    # Check first row (header)
    if table.rows:
        first_row = table.rows[0]
        print("  Header row cells:")
        for cell_idx, cell in enumerate(first_row.cells[:3]):  # First 3 cells
            tc = cell._element
            tcPr = tc.find('{http://schemas.openxmlformats.org/wordprocessingml/2006/main}tcPr')
            if tcPr is not None:
                shd = tcPr.find('{http://schemas.openxmlformats.org/wordprocessingml/2006/main}shd')
                if shd is not None:
                    fill_color = shd.get('{http://schemas.openxmlformats.org/wordprocessingml/2006/main}fill')
                    print(f"    Cell {cell_idx + 1}: Background = {fill_color}")

        # Check second row (data) for NO background
        if len(table.rows) > 1:
            second_row = table.rows[1]
            print("  Data row cells (should have NO background):")
            has_background = False
            for cell_idx, cell in enumerate(second_row.cells[:3]):
                tc = cell._element
                tcPr = tc.find('{http://schemas.openxmlformats.org/wordprocessingml/2006/main}tcPr')
                if tcPr is not None:
                    shd = tcPr.find('{http://schemas.openxmlformats.org/wordprocessingml/2006/main}shd')
                    if shd is not None:
                        fill_color = shd.get('{http://schemas.openxmlformats.org/wordprocessingml/2006/main}fill')
                        if fill_color and fill_color != 'auto':
                            has_background = True
                            print(f"    Cell {cell_idx + 1}: Background = {fill_color} (UNEXPECTED)")
            if not has_background:
                print("    ✓ No background colors (correct)")

# 3. Check signature table structure
print("\n3. SIGNATURE TABLE STRUCTURE:")
print("-" * 80)

last_table = doc.tables[-1]
print(f"Dimensions: {len(last_table.rows)} rows × {len(last_table.columns)} cols")
print("\nRow structure:")
for idx, row in enumerate(last_table.rows):
    cell_text = row.cells[0].text[:30]
    print(f"  Row {idx + 1}: {cell_text}")

# Check signature row height
sig_row = last_table.rows[1]  # Signature row
tr = sig_row._element
trPr = tr.find('{http://schemas.openxmlformats.org/wordprocessingml/2006/main}trPr')
if trPr is not None:
    trHeight = trPr.find('{http://schemas.openxmlformats.org/wordprocessingml/2006/main}trHeight')
    if trHeight is not None:
        height_val = trHeight.get('{http://schemas.openxmlformats.org/wordprocessingml/2006/main}val')
        print(f"\nSignature row height: {height_val} (should be 720 for 0.5\")")

# 4. Check TOC spacing
print("\n4. TABLE OF CONTENTS SPACING:")
print("-" * 80)

toc_found = False
for idx, para in enumerate(doc.paragraphs):
    if 'Table of Contents' in para.text and para.style.name == 'Heading 1':
        toc_found = True
        # Check next few paragraphs
        for i in range(idx + 1, min(idx + 4, len(doc.paragraphs))):
            toc_para = doc.paragraphs[i]
            if toc_para.text.strip():
                space_before = toc_para.paragraph_format.space_before
                space_after = toc_para.paragraph_format.space_after
                print(f"TOC entry: '{toc_para.text[:40]}'")
                print(f"  Space before: {space_before}")
                print(f"  Space after: {space_after}")
                break
        break

print("\n" + "=" * 80)
print("VERIFICATION COMPLETE")
print("=" * 80)
